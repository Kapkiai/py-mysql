from docx import Document
import json

document = Document('871919 - RIVERSIDE HOSTELS SERVICES-2-converted.docx')
tables = document.tables

tab = []

for u in tables:
    h = [[row.cells[a].paragraphs[0].text for row in u.columns]
         for a in range(len(u.rows))]
    dic = {x: list(y) for (x, y) in zip(h[0], zip(*h[1:]))}
    keys = ['Receipt No.', 'Completion Time', 'Details', 'Paid In', 'Withdrawn', 'Balance']
    tab.append({key: dic[key] for key in dic if key in keys})


res = dict()
for dict in tab:
    for list in dict:
        if list in res:
            res[list] += (dict[list])
        else:
            res[list] = dict[list]

with open('test.json', 'w') as fd:
    json.dump(res, fd)
